import pandas as pd
import faker
import random
from datetime import datetime
import json
from docx import Document
from fpdf import FPDF

class DummyDataGenerator:
    def __init__(self, sensitivity):
        # Initialize the DummyDataGenerator object
        self.sensitivity = str(sensitivity)  # Convert sensitivity to a string
        self.fake = faker.Faker("en_AU")  # Create a Faker object for generating fake data
        self.sensitivity_map = {
            "1": "Low",
            "2": "Medium",
            "3": "High",
            "4": "Mixed"
        }
        self.sensitivity_label = self.sensitivity_map.get(self.sensitivity, "Unknown")  # Get the sensitivity label based on the sensitivity value

    def get_sensitive_data(self):
        # Check if sensitive data should be generated
        if self.fake.boolean():
            # Check the sensitivity label and generate corresponding sensitive data
            if self.sensitivity_label == "Low":
                return f"Name: {self.fake.name()}, Email: {self.fake.email()}"
            elif self.sensitivity_label == "Medium":
                return f"Phone: {self.fake.phone_number()}, Address: {self.fake.address()}"
            elif self.sensitivity_label == "High":
                return f"SSN: {self.fake.ssn()}, Credit Card: {self.fake.credit_card_number()}"
            elif self.sensitivity_label == "Mixed":
                return (f"Name: {self.fake.name()}, Phone: {self.fake.phone_number()}, "
                        f"Address: {self.fake.address()}, SSN: {self.fake.ssn()}")
        # Return None if no sensitive data should be generated
        return None

    def generate_csv(self):
        data = []
        num_rows = random.randint(50, 500)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"dummy_data_{self.sensitivity_label}_{timestamp}.csv"

        for _ in range(num_rows):
            row = {
                "Name": self.fake.name(),  # Generate a fake name
                "Email": self.fake.email()  # Generate a fake email
            }
            if self.sensitivity_label in ["Medium", "High", "Mixed"]:
                row["Phone"] = self.fake.phone_number()  # Add a fake phone number
                row["Address"] = self.fake.address()  # Add a fake address
            if self.sensitivity_label in ["High", "Mixed"]:
                row["Date of Birth"] = self.fake.date_of_birth(minimum_age=18, maximum_age=90)  # Add a fake date of birth
                row["SSN"] = self.fake.ssn()  # Add a fake SSN

            data.append(row)

        pd.DataFrame(data).to_csv(file_name, index=False)  # Save data to a CSV file
        print(f"✅ CSV generated: {file_name}")  # Print a success message with the file name

    def generate_json(self, num_records):
        # Generate a timestamp for the file name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Create a file name using sensitivity label and timestamp
        file_name = f"dummy_data_{self.sensitivity_label}_{timestamp}.json"

        data = []  # Initialize a list to store all generated profiles

        for _ in range(num_records):
            # Create a profile with basic user information
            profile = {
                "user_id": self.fake.uuid4(),  # Generate a unique user ID
                "profile": {
                    "name": self.fake.name(),  # Generate a fake name
                    "email": self.fake.email()  # Generate a fake email
                },
                "preferences": {
                    "language": self.fake.language_name(),  # Generate a fake language
                    "timezone": "Australia/Melbourne",  # Set a fixed timezone
                    "marketing_opt_in": self.fake.boolean()  # Random boolean for marketing opt-in
                },
                "account": {
                    "created_at": self.fake.date_time_between(start_date='-5y', end_date='now').isoformat(),  # Random account creation date
                    "last_login": self.fake.date_time_between(start_date='-30d', end_date='now').isoformat(),  # Random last login date
                    "account_status": random.choice(["active", "suspended", "closed"])  # Random account status
                }
            }

            # Add additional sensitive data based on the sensitivity label
            if self.sensitivity_label in ["Medium", "High", "Mixed"]:
                profile["profile"]["phone"] = self.fake.phone_number()  # Add a fake phone number
                profile["profile"]["address"] = self.fake.address()  # Add a fake address
            if self.sensitivity_label in ["High", "Mixed"]:
                profile["profile"]["dob"] = str(self.fake.date_of_birth(minimum_age=18, maximum_age=90))  # Add a fake date of birth
                profile["account"]["api_key"] = f"AKIA{self.fake.random_number(digits=16, fix_len=True)}"  # Add a fake API key

            # Append the generated profile to the data list
            data.append(profile)

        # Write the generated data to a JSON file
        with open(file_name, "w") as file:
            json.dump(data, file, indent=4)  # Save data with pretty formatting

        # Print a success message with the file name
        print(f"✅ JSON file generated: {file_name}")

    def generate_txt(self):
        # Generate a random number of paragraphs between 5 and 20
        num_paragraphs = random.randint(5, 20)
        lines = []  # Initialize a list to store lines of text

        for _ in range(num_paragraphs):
            # Add a paragraph with fake data (20 sentences) to the lines list
            lines.append(self.fake.paragraph(nb_sentences=20))
            
            # Retrieve sensitive data and add it to the lines list if it exists
            sensitive = self.get_sensitive_data()
            if sensitive:
                lines.append(sensitive)

        # Generate a timestamp for the file name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Create a file name using sensitivity level and timestamp
        file_name = f"generated_txt_{self.sensitivity}_{timestamp}.txt"
        
        # Write all lines to the text file
        with open(file_name, "w") as file:
            file.write("\n\n".join(lines))  # Separate paragraphs with double newlines

        # Print a success message with the file name
        print(f"✅ Text file generated: {file_name}")

    def generate_docx(self):
        # Generate a random number of paragraphs between 5 and 20
        num_paragraphs = random.randint(5, 20)
        doc = Document()  # Create a new Word document

        for _ in range(num_paragraphs):
            # Add a paragraph with fake data (20 sentences)
            doc.add_paragraph(self.fake.paragraph(nb_sentences=20))
            
            # Retrieve sensitive data and add it as a separate paragraph if it exists
            sensitive = self.get_sensitive_data()
            if sensitive:
                doc.add_paragraph(sensitive)

        # Generate a timestamp for the file name
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Create a file name using sensitivity level and timestamp
        file_name = f"generated_docx_{self.sensitivity}_{timestamp}.docx"
        # Save the Word document to the file
        doc.save(file_name)

        # Print a success message with the file name
        print(f"✅ Word document generated: {file_name}")

    def generate_pdf(self):
        num_paragraphs = random.randint(5, 20)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        page_width = pdf.w - 2 * pdf.l_margin  # Calculate usable page width

        for _ in range(num_paragraphs):
            paragraph = self.fake.paragraph(nb_sentences=20)
            pdf.set_x(pdf.l_margin)  # Reset X position to the left margin
            pdf.multi_cell(page_width, 10, paragraph, align="L")  # Align text to the left
            sensitive = self.get_sensitive_data()
            if isinstance(sensitive, str) and sensitive.strip():
                pdf.set_x(pdf.l_margin)  # Reset X position for sensitive data
                pdf.multi_cell(page_width, 10, sensitive, align="L")  # Align sensitive data to the left
            pdf.ln(5)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_file = f"generated_pdf_{self.sensitivity}_{timestamp}.pdf"
        pdf.output(pdf_file)

        print(f"✅ PDF generated: {pdf_file}")